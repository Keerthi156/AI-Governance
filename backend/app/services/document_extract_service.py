"""
Extract plain text from uploaded knowledge files for RAG ingest.

Supported: .txt, .md, .markdown, .pdf, .docx
"""

from __future__ import annotations

import io
from pathlib import Path

from app.core.exceptions import ValidationAppError

ALLOWED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}
MAX_UPLOAD_BYTES = 2 * 1024 * 1024  # 2 MiB
MAX_CONTENT_CHARS = 200_000


def _decode_text_bytes(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValidationAppError("Could not decode text file as UTF-8 or Latin-1")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ValidationAppError(
            "PDF support requires the pypdf package on the server"
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ValidationAppError(f"Invalid or corrupt PDF: {exc}") from exc

    parts: list[str] = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        if text.strip():
            parts.append(text.strip())
    content = "\n\n".join(parts).strip()
    if not content:
        raise ValidationAppError("PDF contained no extractable text")
    return content


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValidationAppError(
            "DOCX support requires the python-docx package on the server"
        ) from exc

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ValidationAppError(f"Invalid or corrupt DOCX: {exc}") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    # Include simple table cell text (common in policy handbooks).
    for table in document.tables:
        for row in table.rows:
            cells = [
                (cell.text or "").strip()
                for cell in row.cells
                if (cell.text or "").strip()
            ]
            if cells:
                parts.append(" | ".join(cells))

    content = "\n\n".join(parts).strip()
    if not content:
        raise ValidationAppError("DOCX contained no extractable text")
    return content


def extract_text_from_upload(*, filename: str, data: bytes) -> tuple[str, str]:
    """
    Return (extracted_text, normalized_extension).

    Raises ValidationAppError for unsupported types or empty content.
    """
    if not filename or not filename.strip():
        raise ValidationAppError("Upload filename is required")
    if not data:
        raise ValidationAppError("Uploaded file is empty")
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValidationAppError(
            f"File exceeds maximum size of {MAX_UPLOAD_BYTES // (1024 * 1024)} MiB"
        )

    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValidationAppError(
            f"Unsupported file type “{ext or '(none)'}”. "
            f"Allowed: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    else:
        text = _decode_text_bytes(data).strip()

    if not text:
        raise ValidationAppError("Extracted document text is empty")
    if len(text) > MAX_CONTENT_CHARS:
        raise ValidationAppError(
            f"Extracted text exceeds {MAX_CONTENT_CHARS:,} characters"
        )
    return text, ext
